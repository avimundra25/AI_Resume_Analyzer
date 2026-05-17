"""
Resume Parser - Production-Ready Implementation
Uses pdfplumber, python-docx, spaCy, and sentence-transformers for intelligent resume parsing.
"""

import re
from pathlib import Path
from typing import Optional

import pdfplumber
from docx import Document
import spacy
from sentence_transformers import SentenceTransformer, util

# Load NLP models (cached after first load)
nlp = spacy.load("en_core_web_sm")
model = SentenceTransformer('all-MiniLM-L6-v2')

# Skill database with common variations
SKILL_DB = [
    "python", "java", "c++", "javascript", "typescript",
    "react", "node.js", "angular", "vue.js",
    "sql", "mysql", "postgresql", "mongodb", "html", "css",
    "machine learning", "deep learning", "natural language processing",
    "data science", "data analysis", "statistics",
    "flask", "django", "fastapi", "spring boot",
    "docker", "kubernetes", "aws", "azure", "gcp",
    "git", "linux", "rest api", "graphql",
    "communication", "teamwork", "leadership", "problem solving",
    "critical thinking", "project management", "agile", "scrum",
    "cleaning equipment", "warehouse operations", "inventory management",
    "logistics", "customer service", "sales",
    "patient care", "diagnosis", "healthcare", "medical records",
    "accounting", "financial analysis", "excel", "powerpoint"
]

# Skill normalization mapping
SKILL_ALIASES = {
    "ml": "machine learning",
    "dl": "deep learning",
    "nlp": "natural language processing",
    "js": "javascript",
    "ts": "typescript",
    "py": "python",
    "postgres": "postgresql",
    "mongo": "mongodb",
    "k8s": "kubernetes",
    "ec2": "aws",
    "s3": "aws",
}

# Role database with descriptive text for semantic matching
ROLE_DB = {
    "data scientist": "machine learning, statistics, data analysis, python, deep learning, predictive modeling, data visualization, pandas, numpy, scikit-learn",
    "software engineer": "programming, software development, system design, algorithms, coding, debugging, version control, object-oriented programming",
    "frontend developer": "html, css, javascript, react, angular, vue, user interface, responsive design, web development",
    "backend developer": "server-side programming, databases, api development, node.js, python, java, system architecture",
    "full stack developer": "frontend, backend, databases, javascript, react, node.js, full application development",
    "data analyst": "data analysis, sql, excel, visualization, reporting, business intelligence, tableau, power bi",
    "devops engineer": "ci/cd, docker, kubernetes, cloud infrastructure, automation, linux, aws, azure",
    "product manager": "product strategy, roadmap planning, stakeholder management, agile, user research, requirements gathering",
    "doctor": "patient care, diagnosis, treatment, medical knowledge, healthcare, clinical skills, prescription",
    "nurse": "patient care, medical assistance, healthcare, vital signs monitoring, medication administration",
    "accountant": "accounting, financial statements, tax preparation, bookkeeping, auditing, excel",
    "marketing manager": "digital marketing, brand management, campaign strategy, social media, content marketing, analytics",
    "hr manager": "recruitment, employee relations, performance management, training, compliance, organizational development",
    "warehouse associate": "inventory management, logistics, shipping, receiving, forklift operation, warehouse operations",
    "customer service representative": "customer support, communication, problem resolution, phone support, crm systems"
}

# Section header patterns
SECTION_PATTERNS = {
    "skills": r"(?i)^(skills?|technical\s+skills?|core\s+competenc|expertise|proficienc)",
    "education": r"(?i)^(education|academic|qualification|degree)",
    "experience": r"(?i)^(experience|employment|work\s+history|professional\s+experience|career)",
    "projects": r"(?i)^(projects?|portfolio)",
    "certifications": r"(?i)^(certification|certificate|license)",
    "summary": r"(?i)^(summary|objective|profile|about)"
}

# Degree patterns
DEGREE_PATTERNS = [
    r"(?i)\b(ph\.?d\.?|doctorate)\b",
    r"(?i)\b(m\.?b\.?a\.?)\b",
    r"(?i)\b(m\.?s\.?|master(?:'?s)?(?:\s+of\s+science)?)\b",
    r"(?i)\b(m\.?tech\.?|m\.?\s*tech)\b",
    r"(?i)\b(b\.?tech\.?|b\.?\s*tech)\b",
    r"(?i)\b(b\.?e\.?|bachelor(?:'?s)?(?:\s+of\s+engineering)?)\b",
    r"(?i)\b(b\.?s\.?|b\.?sc\.?|bachelor(?:'?s)?(?:\s+of\s+science)?)\b",
    r"(?i)\b(b\.?a\.?|bachelor(?:'?s)?(?:\s+of\s+arts)?)\b",
    r"(?i)\b(b\.?com\.?|bachelor(?:'?s)?(?:\s+of\s+commerce)?)\b",
    r"(?i)\b(m\.?b\.?b\.?s\.?|mbbs)\b",
    r"(?i)\b(b\.?c\.?a\.?|bca)\b",
    r"(?i)\b(m\.?c\.?a\.?|mca)\b",
    r"(?i)\b(diploma)\b",
    r"(?i)\b(high\s+school|12th|hsc|intermediate)\b",
]

# Year pattern
YEAR_PATTERN = r"\b(19|20)\d{2}\b"

# Duration patterns for experience
DURATION_PATTERNS = [
    r"(\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*\d{4})\s*[-–—to]+\s*(present|\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*\d{4})",
    r"(\d{1,2}/\d{4})\s*[-–—to]+\s*(present|\d{1,2}/\d{4})",
    r"(\d{4})\s*[-–—to]+\s*(present|\d{4})",
    r"(\d+)\s*(?:years?|yrs?)",
]

# Precompute embeddings
skill_embeddings = model.encode(SKILL_DB, convert_to_tensor=True)
role_descriptions = list(ROLE_DB.values())
role_names = list(ROLE_DB.keys())
role_embeddings = model.encode(role_descriptions, convert_to_tensor=True)


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text content from a PDF file.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text as a string
        
    Raises:
        ValueError: If the PDF cannot be read or is empty
    """
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            if not pdf.pages:
                raise ValueError("PDF has no pages")
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        raise ValueError(f"Failed to read PDF: {e}")
    
    if not text.strip():
        raise ValueError("PDF contains no extractable text")
    
    return text


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text content from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text as a string
        
    Raises:
        ValueError: If the DOCX cannot be read or is empty
    """
    try:
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        text = "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"Failed to read DOCX: {e}")
    
    if not text.strip():
        raise ValueError("DOCX contains no extractable text")
    
    return text


def extract_text(file_path: str) -> str:
    """
    Extract text from PDF or DOCX file based on extension.
    
    Args:
        file_path: Path to the resume file
        
    Returns:
        Extracted text as a string
        
    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If file type is unsupported or file is unreadable
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    extension = path.suffix.lower()
    
    if extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif extension in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {extension}. Use PDF or DOCX.")


def detect_sections(text: str) -> dict[str, str]:
    """
    Detect and extract content by section headers.
    
    Args:
        text: Full resume text
        
    Returns:
        Dictionary mapping section names to their content
    """
    lines = text.split('\n')
    sections = {}
    current_section = "header"
    current_content = []
    
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            current_content.append(line)
            continue
        
        # Check if line matches any section header
        matched_section = None
        for section_name, pattern in SECTION_PATTERNS.items():
            if re.match(pattern, line_stripped):
                matched_section = section_name
                break
        
        if matched_section:
            # Save previous section
            if current_content:
                sections[current_section] = '\n'.join(current_content)
            current_section = matched_section
            current_content = []
        else:
            current_content.append(line)
    
    # Save last section
    if current_content:
        sections[current_section] = '\n'.join(current_content)
    
    return sections


def extract_email(text: str) -> Optional[str]:
    """
    Extract email address from text.
    
    Args:
        text: Text to search
        
    Returns:
        Email address or None if not found
    """
    # More robust email pattern
    pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
    match = re.search(pattern, text)
    return match.group().lower() if match else None


def extract_phone(text: str) -> Optional[str]:
    """
    Extract phone number from text.
    
    Args:
        text: Text to search
        
    Returns:
        Phone number or None if not found
    """
    # Multiple phone patterns
    patterns = [
        r"\+?1?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}",  # US format
        r"\+?\d{2,3}[-.\s]?\d{10}",  # International with country code
        r"\b\d{10}\b",  # Simple 10 digits
        r"\b\d{5}[-.\s]?\d{5}\b",  # 5-5 format
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            # Clean up the number
            phone = re.sub(r"[^\d+]", "", match.group())
            if len(phone) >= 10:
                return phone
    
    return None


def extract_name(text: str) -> Optional[str]:
    """
    Extract candidate name using NLP entity recognition.
    
    Args:
        text: Resume text (uses first 500 chars for efficiency)
        
    Returns:
        Candidate name or None if not found
    """
    # Focus on the header section where name typically appears
    header_text = text[:500]
    doc = nlp(header_text)
    
    # First, try to find PERSON entities
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            name = ent.text.strip()
            # Basic validation: name should have at least 2 characters
            # and shouldn't be common false positives
            if len(name) > 2 and not re.match(r"^(resume|cv|curriculum)$", name.lower()):
                return name
    
    # Fallback: first line often contains the name
    first_lines = text.strip().split('\n')[:3]
    for line in first_lines:
        line = line.strip()
        # Check if it looks like a name (2-4 words, no special chars)
        if line and 2 <= len(line.split()) <= 4:
            if re.match(r"^[A-Za-z\s.'-]+$", line):
                return line
    
    return None


def normalize_skill(skill: str) -> str:
    """
    Normalize skill text using alias mapping.
    
    Args:
        skill: Raw skill text
        
    Returns:
        Normalized skill name
    """
    skill_lower = skill.lower().strip()
    return SKILL_ALIASES.get(skill_lower, skill_lower)


def extract_skills(text: str, sections: dict[str, str]) -> list[str]:
    """
    Improved skill extraction using controlled candidates + semantic filtering
    """

    # ✅ 1. PRIORITIZE SKILLS + PROJECTS SECTION ONLY
    skills_text = sections.get("skills", "") + "\n" + sections.get("projects", "")
    
    # fallback if no sections found
    if not skills_text.strip():
        skills_text = text[:2000]  # only first part of resume (important)

    doc = nlp(skills_text)

    candidates = set()

    
    for chunk in doc.noun_chunks:
        phrase = chunk.text.strip().lower()
        words = phrase.split()

        if (
            1 <= len(words) <= 3 and
            len(phrase) > 2 and
            not re.match(r"^\d+$", phrase)
        ):
            candidates.add(normalize_skill(phrase))

    

    
    text_lower = skills_text.lower()
    for skill in SKILL_DB:
        if skill in text_lower:
            candidates.add(skill)

    if not candidates:
        return []

    candidates = list(candidates)

    # ✅ 4. SEMANTIC MATCHING (STRONGER FILTER)
    candidate_embeddings = model.encode(candidates, convert_to_tensor=True)
    similarities = util.cos_sim(candidate_embeddings, skill_embeddings)

    found_skills = set()
    skill_scores = {}

    for i, candidate in enumerate(candidates):
        for j, skill in enumerate(SKILL_DB):
            score = similarities[i][j].item()

            # 🔥 HIGHER THRESHOLD (important fix)
            if score > 0.70:
                if skill not in skill_scores or score > skill_scores[skill]:
                    skill_scores[skill] = score
                    found_skills.add(skill)

    # ✅ 5. SORT + LIMIT
    sorted_skills = sorted(
        found_skills,
        key=lambda s: skill_scores.get(s, 0),
        reverse=True
    )

    return sorted_skills[:12]  # slightly tighter list
def extract_education(text: str, sections: dict[str, str]) -> dict[str, Optional[str]]:
    """
    Extract education information including degree, college, and year.
    
    Args:
        text: Full resume text
        sections: Detected sections dictionary
        
    Returns:
        Dictionary with degree, college, and year
    """
    education_text = sections.get("education", text)
    
    result = {
        "degree": None,
        "college": None,
        "year": None
    }
    
    # Extract degree
    for pattern in DEGREE_PATTERNS:
        match = re.search(pattern, education_text)
        if match:
            result["degree"] = match.group().strip()
            break
    
    # Extract year (prefer graduation year - usually the later one in education section)
    years = re.findall(YEAR_PATTERN, education_text)
    if years:
        # Get the most recent year that's not in the future
        from datetime import datetime
        current_year = datetime.now().year
        valid_years = [y for y in years if int(y) <= current_year + 5]
        if valid_years:
            result["year"] = max(valid_years)
    
    # Extract college/university using NLP
    doc = nlp(education_text)
    for ent in doc.ents:
        if ent.label_ == "ORG":
            org_text = ent.text.strip()
            # Check if it looks like an educational institution
            if re.search(r"(?i)(university|college|institute|school|academy)", org_text):
                result["college"] = org_text
                break
    
    # Fallback: look for common university patterns
    if not result["college"]:
        uni_pattern = r"(?i)([\w\s]+(?:university|college|institute|school|academy)[\w\s]*)"
        match = re.search(uni_pattern, education_text)
        if match:
            result["college"] = match.group().strip()
    
    return result


def extract_experience(text: str, sections: dict[str, str]) -> list[dict[str, Optional[str]]]:
    """
    Extract work experience including role, company, and duration.
    
    Args:
        text: Full resume text
        sections: Detected sections dictionary
        
    Returns:
        List of experience dictionaries
    """
    experience_text = sections.get("experience", "")
    if not experience_text:
        return []
    
    experiences = []
    doc = nlp(experience_text)
    
    # Split into potential job entries (by blank lines or bullet points)
    entries = re.split(r'\n\s*\n|\n(?=•|▪|►|■|\*|-\s)', experience_text)
    
    for entry in entries:
        if len(entry.strip()) < 10:
            continue
        
        exp = {
            "role": None,
            "company": None,
            "duration": None
        }
        
        entry_doc = nlp(entry)
        
        # Extract company (ORG entities)
        for ent in entry_doc.ents:
            if ent.label_ == "ORG" and not exp["company"]:
                exp["company"] = ent.text.strip()
        
        # Extract duration
        for pattern in DURATION_PATTERNS:
            match = re.search(pattern, entry, re.IGNORECASE)
            if match:
                exp["duration"] = match.group().strip()
                break
        
        # Extract role (common job title patterns)
        role_patterns = [
            r"(?i)(senior|junior|lead|chief|head|principal)?\s*(software|data|product|project|marketing|sales|hr|finance|operations)?\s*(engineer|developer|analyst|scientist|manager|director|executive|consultant|specialist|coordinator|associate|intern)",
            r"(?i)(ceo|cto|cfo|coo|vp|svp|evp)",
            r"(?i)(intern|trainee|apprentice)",
        ]
        
        for pattern in role_patterns:
            match = re.search(pattern, entry)
            if match:
                exp["role"] = match.group().strip()
                break
        
        # Only add if we found at least one field
        if exp["role"] or exp["company"] or exp["duration"]:
            experiences.append(exp)
    
    return experiences[:5]  # Limit to 5 most recent


def predict_role(skills: list[str]) -> tuple[Optional[str], float]:
    """
    Predict the most suitable role based on extracted skills using semantic matching.
    
    Args:
        skills: List of extracted skills
        
    Returns:
        Tuple of (predicted role name, confidence score)
    """
    if not skills:
        return None, 0.0
    
    # Create a skill profile from extracted skills
    skill_profile = ", ".join(skills)
    profile_embedding = model.encode(skill_profile, convert_to_tensor=True)
    
    # Compare with role descriptions
    similarities = util.cos_sim(profile_embedding, role_embeddings)
    
    # Get the best match
    best_idx = similarities.argmax().item()
    confidence = similarities[0][best_idx].item()
    
    # Only return if confidence is reasonable
    if confidence > 0.3:
        return role_names[best_idx], round(confidence, 2)
    
    return None, 0.0


def parse_resume(file_path: str) -> dict:
    """
    Parse a resume file and extract structured information.
    
    Main entry point for the resume parser.
    
    Args:
        file_path: Path to the resume file (PDF or DOCX)
        
    Returns:
        Dictionary containing:
        - name: Candidate name
        - email: Email address
        - phone: Phone number
        - skills: List of identified skills
        - education: Education details (degree, college, year)
        - experience: List of work experiences
        - predicted_role: AI-predicted suitable role
        - confidence: Confidence score for role prediction
        
    Raises:
        FileNotFoundError: If the file doesn't exist
        ValueError: If the file type is unsupported or unreadable
    """
    # Extract text from file
    text = extract_text(file_path)
    
    # Detect sections for targeted extraction
    sections = detect_sections(text)
    
    # Extract all fields
    skills = extract_skills(text, sections)
    predicted_role, confidence = predict_role(skills)
    
    return {
        "name": extract_name(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
        "skills": skills,
        "education": extract_education(text, sections),
        "experience": extract_experience(text, sections),
        "predicted_role": predicted_role,
        "confidence": confidence
    }


# Example usage
if __name__ == "__main__":
    import sys
    import json
    
    if len(sys.argv) < 2:
        print("Usage: python parser.py <resume_file>")
        sys.exit(1)
    
    try:
        result = parse_resume(sys.argv[1])
        print(json.dumps(result, indent=2))
    except (FileNotFoundError, ValueError) as e:
        print(f"Error: {e}")
        sys.exit(1)
