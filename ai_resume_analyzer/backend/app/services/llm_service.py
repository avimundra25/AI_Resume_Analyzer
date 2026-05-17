import json
import logging
from typing import Dict, Optional
import google.generativeai as genai
from docx import Document
import os
from app.core.config import settings

logger = logging.getLogger(__name__)

if not settings.GEMINI_API_KEY:
    logger.error("GEMINI_API_KEY not found! Make sure your .env file is set up.")
else:
    genai.configure(api_key=settings.GEMINI_API_KEY)

def generate_evaluation_report(parsed_resume_data: Dict, job_description: str, output_filename: str = "AI_Recommendations_Report.docx") -> dict:
    logger.info("Initializing Gemini API call...")
    model = genai.GenerativeModel('gemini-1.5-flash')
    resume_string = json.dumps(parsed_resume_data, indent=2)
    
    # UPDATED PROMPT: Forces the LLM to output the EXACT JSON format the frontend UI expects
    prompt = f"""
    Act as an expert technical recruiter. Analyze this structured Resume Data against the Job Description.
    Provide the output STRICTLY as a valid JSON object with NO markdown formatting, NO backticks, and exactly these keys:
    "score": <integer 0-100>,
    "interview_probability": <integer 0-100>,
    "overall_summary": "<2-3 sentence technical summary>",
    "strengths": ["<strength 1>", "<strength 2>", "<strength 3>"],
    "weaknesses": ["<weakness 1>", "<weakness 2>"],
    "suggestions": ["<actionable tip 1>", "<actionable tip 2>"],
    "missing_skills": ["<skill 1>", "<skill 2>"],
    "content_quality": <integer 0-100>,
    "skills_match": <integer 0-100>,
    "structure_format": <integer 0-100>,
    "ats_compatibility": <integer 0-100>

    Resume Data:
    {resume_string}

    Job Description:
    {job_description}
    """

    try:
        response = model.generate_content(prompt)
        clean_text = response.text.strip().replace('```json', '').replace('```', '').strip()
        analysis_data = json.loads(clean_text)
        
        logger.info("LLM Evaluation successful. Compiling Word Document...")

        # Generate the Word Document using the new data structure
        doc = Document()
        doc.add_heading('AI Resume Analysis & Recommendation Report', 0)

        candidate_name = parsed_resume_data.get("name", "Candidate")
        doc.add_paragraph(f"Prepared for: {candidate_name}")
        doc.add_paragraph("_" * 50) 

        doc.add_heading('1. Suitability Metrics', level=1)
        doc.add_paragraph(f"Overall ATS Score: {analysis_data.get('score', 'N/A')}/100")
        doc.add_paragraph(f"Interview Probability: {analysis_data.get('interview_probability', 'N/A')}%")
        doc.add_paragraph(f"Skills Match: {analysis_data.get('skills_match', 'N/A')}%")

        doc.add_heading('2. Recruiter Summary', level=1)
        doc.add_paragraph(analysis_data.get('overall_summary', 'No feedback provided.'))

        doc.add_heading('3. Key Strengths', level=1)
        for point in analysis_data.get('strengths', []):
            doc.add_paragraph(point, style='List Bullet')
            
        doc.add_heading('4. Areas for Improvement', level=1)
        for point in analysis_data.get('suggestions', []):
            doc.add_paragraph(point, style='List Bullet')

        full_output_path = os.path.join(settings.REPORT_DIR, output_filename)
        doc.save(full_output_path)
        
        return {
            "status": "success",
            "report_path": full_output_path,
            "data": analysis_data 
        }

    except Exception as e:
        logger.error(f"Failed to generate evaluation: {e}")
        raise ValueError(f"AI Evaluation Failed: {e}")